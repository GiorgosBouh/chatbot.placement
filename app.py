import streamlit as st
import json
import re
import os
import datetime
import requests
import io
from typing import List, Dict, Tuple
from dataclasses import dataclass

# Import Groq with fallback handling
try:
    from groq import Groq
    GROQ_AVAILABLE = True
    print("✅ Groq library imported successfully")
except ImportError:
    GROQ_AVAILABLE = False
    Groq = None
    print("⚠️ Groq library not available. Using fallback mode only.")

# Import python-docx with fallback handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx library imported successfully")
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    print("⚠️ python-docx library not available. DOCX search disabled.")

# Ρύθμιση σελίδας
st.set_page_config(
    page_title="Πρακτική Άσκηση - Μητροπολιτικό Κολλέγιο",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="collapsed"
)

@dataclass
class QAEntry:
    id: int
    category: str
    question: str
    answer: str
    keywords: List[str]

class InternshipChatbot:
    def __init__(self, groq_api_key: str = None):
        # Initialize Groq client if available and API key provided
        self.groq_client = None
        if GROQ_AVAILABLE and groq_api_key:
            try:
                self.groq_client = Groq(api_key=groq_api_key)
                print("✅ Groq client initialized")
            except Exception as e:
                print(f"⚠️ Failed to initialize Groq: {e}")
        
        # Load Q&A data
        self.qa_data = self.load_qa_data()
        
        # Initialize DOCX files cache
        self.docx_cache = {}
        self.docx_files = [
            "1.ΑΙΤΗΣΗ ΠΡΑΓΜΑΤΟΠΟΙΗΣΗΣ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ.DOCX",
            "2.ΣΤΟΙΧΕΙΑ ΔΟΜΗΣ_ΟΔΗΓΙΕΣ.docx", 
            "3.ΣΤΟΙΧΕΙΑ ΦΟΙΤΗΤΗ.docx",
            "4.ΣΤΟΙΧΕΙΑ ΦΟΡΕΑ.docx",
            "5.ΑΣΦΑΛΙΣΤΙΚΗ ΙΚΑΝΟΤΗΤΑ.docx",
            "6.ΥΠΕΥΘΥΝΗ ΔΗΛΩΣΗ Ν 105-Πρακτικής.docx",
            "8.ΒΙΒΛΙΟ_ΠΡΑΚΤΙΚΗΣ_final.docx"
        ]
        
        # System prompt για το LM
        self.system_prompt = """Είσαι ένας εξειδικευμένος σύμβουλος για θέματα πρακτικής άσκησης στο Μητροπολιτικό Κολλέγιο Θεσσαλονίκης, τμήμα Προπονητικής και Φυσικής Αγωγής.

ΚΡΙΣΙΜΕΣ ΓΛΩΣΣΙΚΕΣ ΟΔΗΓΙΕΣ:
- Χρησιμοποίησε ΑΠΟΚΛΕΙΣΤΙΚΑ και ΜΟΝΟ ελληνικούς χαρακτήρες
- ΑΠΑΓΟΡΕΥΟΝΤΑΙ: αγγλικά, κινέζικα, greeklish ή οποιοιδήποτε άλλοι χαρακτήρες
- Ελέγχισε κάθε λέξη πριν την εκτύπωση - πρέπει να είναι ελληνική
- Αν δεν ξέρεις ελληνική λέξη, χρησιμοποίησε περιφραστικό τρόπο

ΚΡΙΤΙΚΕΣ ΟΔΗΓΙΕΣ:
- Μην προσθέτεις πληροφορίες που δεν υπάρχουν στο context
- Χρησιμοποίησε ΜΟΝΟ τις πληροφορίες που σου δίνονται
- Μην εφευρίσκεις ή μην υποθέτεις στοιχεία

ΣΤΥΛ ΑΠΑΝΤΗΣΗΣ:
- Αυστηρά επίσημος και επαγγελματικός τόνος
- Άμεσες και συγκεκριμένες οδηγίες
- Χωρίς χαιρετισμούς, φιλικές εκφράσεις ή περιττά λόγια
- Δομημένες απαντήσεις με σαφή βήματα
- Χωρίς emojis ή άτυπες εκφράσεις

ΒΑΣΙΚΕΣ ΠΛΗΡΟΦΟΡΙΕΣ (μόνο αυτές):
- Υπεύθυνος Πρακτικής Άσκησης: Γεώργιος Σοφιανίδης
- Email: gsofianidis@mitropolitiko.edu.gr
- Τεχνική Υποστήριξη: Γεώργιος Μπουχουράς (gbouchouras@mitropolitiko.edu.gr)
- Απαιτούμενες ώρες: 240 ώρες μέχρι 30/5
- Ωράριο: Δευτέρα-Σάββατο, μέχρι 8 ώρες/ημέρα
- Σύμβαση: Ανέβασμα στο moodle μέχρι 15/10

ΤΕΛΙΚΟΣ ΕΛΕΓΧΟΣ:
- Κάθε απάντηση πρέπει να περιέχει ΜΟΝΟ ελληνικούς χαρακτήρες
- Καμία ξένη λέξη ή χαρακτήρας δεν επιτρέπεται
- Επαγγελματικό ύφος χωρίς φιλικότητες

Απάντησε στα ελληνικά με αυστηρά επαγγελματικό τόνο χρησιμοποιώντας μόνο τις δοσμένες πληροφορίες."""

    def load_qa_data(self) -> List[Dict]:
        """Load Q&A data with better error handling and debugging"""
        filename = "qa_data.json"
        
        print(f"🔍 Looking for {filename}...")
        
        # Check if file exists
        if not os.path.exists(filename):
            print(f"❌ File {filename} not found in current directory")
            print(f"📁 Current directory: {os.getcwd()}")
            print(f"📂 Files in directory: {[f for f in os.listdir('.') if f.endswith('.json')]}")
            return self.get_updated_fallback_data()
        
        # Try to load the file
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
            # Validate data structure
            if not isinstance(data, list):
                print(f"❌ Invalid data format in {filename} - expected list")
                return self.get_updated_fallback_data()
            
            if not data:
                print(f"❌ Empty data in {filename}")
                return self.get_updated_fallback_data()
            
            # Check data integrity
            required_fields = ['id', 'category', 'question', 'answer', 'keywords']
            for i, entry in enumerate(data):
                if not all(field in entry for field in required_fields):
                    print(f"❌ Missing fields in entry {i}: {entry.keys()}")
                    return self.get_updated_fallback_data()
            
            print(f"✅ Successfully loaded {len(data)} Q&A entries from {filename}")
            print(f"📊 Entry IDs: {[entry['id'] for entry in data[:5]]}{'...' if len(data) > 5 else ''}")
            return data
            
        except json.JSONDecodeError as e:
            print(f"❌ JSON decode error in {filename}: {e}")
            return self.get_updated_fallback_data()
        except Exception as e:
            print(f"❌ Error loading {filename}: {e}")
            return self.get_updated_fallback_data()

    def get_updated_fallback_data(self) -> List[Dict]:
        """Updated fallback data with more entries"""
        print("📋 Using updated fallback data...")
        return [
            {
                "id": 1,
                "category": "Γενικές Πληροφορίες",
                "question": "Πώς ξεκινάω την πρακτική μου άσκηση;",
                "answer": "1. Επικοινωνώ με τον υπεύθυνο της πρακτικής: gsofianidis@mitropolitiko.edu.gr\n\n2. Βρίσκω τη δομή που θα κάνω πρακτική\n\n3. Κατεβάζω τα έγγραφα από το μάθημα SPORTS COACHING PRACTICE & EXPERTISE DEVELOPMENT (SE5117) στο Moodle. Τα συμπληρώνω και τα ανεβάζω ξανά στη σχετική πύλη στο μάθημα SPORTS COACHING PRACTICE & EXPERTISE DEVELOPMENT (SE5117) στο Moodle.\n\n4. Περιμένω την υπογραφή της σύμβασής μου και την ανάρτησή της στο ΕΡΓΑΝΗ\n\n5. Ξεκινάω την πρακτική",
                "keywords": ["ξεκινάω", "ξεκινώ", "αρχή", "αρχίζω", "αρχίσω", "ξεκίνημα", "πρακτική", "άσκηση", "πώς", "πως", "βήματα", "διαδικασία", "διαδικασιες"]
            },
            {
                "id": 2,
                "category": "Έγγραφα & Διαδικασίες",
                "question": "Τι έγγραφα χρειάζομαι για την πρακτική άσκηση;",
                "answer": "Για εσάς (φοιτητή):\n• Αίτηση πραγματοποίησης πρακτικής άσκησης\n• Στοιχεία φοιτητή (συμπληρωμένη φόρμα)\n• Ασφαλιστική ικανότητα από gov.gr\n• Υπεύθυνη δήλωση (δεν παίρνετε επίδομα ΟΑΕΔ)\n\nΓια τη δομή:\n• Στοιχεία φορέα (ΑΦΜ, διεύθυνση, νόμιμος εκπρόσωπος)\n• Ημέρες και ώρες που σας δέχεται\n\nTip: Ξεκινήστε από την ασφαλιστική ικανότητα γιατί παίρνει χρόνο!",
                "keywords": ["έγγραφα", "εγγραφα", "χαρτιά", "χαρτια", "χρειάζομαι", "χρειαζομαι", "απαιτήσεις", "απαιτησεις", "απαιτούνται", "απαιτουνται", "δικαιολογητικά", "δικαιολογητικα", "φάκελος", "φακελος", "αίτηση", "αιτηση"]
            },
            {
                "id": 30,
                "category": "Οικονομικά & Αμοιβή",
                "question": "Παίρνω αμοιβή για την πρακτική άσκηση; Τι κόστος έχει για τη δομή;",
                "answer": "ΓΙΑ ΤΟΥΣ ΦΟΙΤΗΤΕΣ:\n\nΔΕΝ υπάρχει αμοιβή για την πρακτική άσκηση\n• Η πρακτική άσκηση είναι μη αμειβόμενη\n• Είναι μέρος των σπουδών σας\n• Δεν πρόκειται για εργασιακή σχέση\n\nΓΙΑ ΤΗ ΔΟΜΗ:\n\nΗ δομή δε χρεώνεται κάτι (σχεδόν)\n• Υπάρχει ένα ελάχιστο τέλος που ενδεχομένως πρέπει να καταβάλει\n• Το κολλέγιο καλύπτει τα έξοδα της σύμβασης\n• Η ασφάλιση τιμολογείται στο κολλέγιο\n• Δεν υπάρχει οικονομική υποχρέωση προς τους φοιτητές",
                "keywords": ["αμοιβή", "αμοιβη", "πληρωμή", "πληρωμη", "πληρώθώ", "πληρωθώ", "πληρωθω", "πληρωνομαι", "πληρώνομαι", "λεφτά", "λεφτα", "χρήματα", "χρηματα", "κόστος", "κοστος", "τέλος", "τελος", "δομή", "δομη", "φοιτητής", "φοιτητη", "οικονομικά", "οικονομικα", "μισθός", "μισθος"]
            },
            {
                "id": 11,
                "category": "Επικοινωνία",
                "question": "Με ποιον επικοινωνώ για την πρακτική άσκηση;",
                "answer": "ΚΥΡΙΑ ΕΠΙΚΟΙΝΩΝΙΑ:\n\nΓεώργιος Σοφιανίδης, MSc, PhD(c)\n📧 gsofianidis@mitropolitiko.edu.gr\nΥπεύθυνος Πρακτικής Άσκησης\n\nΕΝΑΛΛΑΚΤΙΚΗ ΕΠΙΚΟΙΝΩΝΙΑ:\n\nΓεώργιος Μπουχουράς, MSc, PhD\n📧 gbouchouras@mitropolitiko.edu.gr\n📞 2314 409000\nProgramme Leader\n\nΠότε να επικοινωνήσετε:\n• Ερωτήσεις για έγγραφα ➜ Γεώργιος Σοφιανίδης\n• Τεχνικά προβλήματα ➜ Γεώργιος Σοφιανίδης\n• Θέματα προγράμματος ➜ Γεώργιος Μπουχουράς",
                "keywords": ["επικοινωνία", "επικοινωνια", "Σοφιανίδης", "Σοφιανιδης", "Μπουχουράς", "Μπουχουρας", "email", "τηλέφωνο", "τηλεφωνο", "υπεύθυνος", "υπευθυνος", "βοήθεια", "βοηθεια", "καθηγητής", "καθηγητης", "καθηγήτρια", "καθηγητρια", "contact", "στοιχεία", "στοιχεια"]
            },
            {
                "id": 4,
                "category": "Ώρες & Χρονοδιάγραμμα",
                "question": "Πόσες ώρες πρέπει να κάνω πρακτική άσκηση;",
                "answer": "Υποχρεωτικό: Τουλάχιστον 240 ώρες\n\nDeadline: Μέχρι 30 Μάϊου\n\nΚανόνες ωραρίου:\n• Δευτέρα έως Σάββατο (ΌΧΙ Κυριακές, 5μέρες/εβδ)\n• Μέχρι 8 ώρες την ημέρα\n• Το ωράριο ορίζεται από τη δομή σε συνεργασία μαζί σας\n\nΥπολογισμός: 240 ώρες = περίπου 6 εβδομάδες x 40 ώρες ή 8 εβδομάδες x 30 ώρες",
                "keywords": ["ώρες", "ωρες", "240", "ποσες", "πόσες", "ποσα", "ποσά", "συνολικά", "συνολικα", "όλες", "ολες", "τελικά", "τελικα", "χρονοδιάγραμμα", "χρονοδιαγραμμα", "διάρκεια", "διαρκεια", "χρόνος", "χρονος", "30/5", "deadline"]
            }
        ]

    def download_docx_file(self, filename: str) -> str:
        """Download and extract text from DOCX file from GitHub"""
        if not DOCX_AVAILABLE:
            print(f"⚠️ python-docx not available, cannot process {filename}")
            return ""
        
        # Check cache first
        if filename in self.docx_cache:
            print(f"📋 Using cached content for {filename}")
            return self.docx_cache[filename]
        
        try:
            # GitHub raw URL
            base_url = "https://raw.githubusercontent.com/GiorgosBouh/chatbot.placement/main/"
            url = base_url + filename
            
            print(f"🔍 Downloading {filename} from GitHub...")
            
            # Download file
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            # Parse DOCX
            doc = Document(io.BytesIO(response.content))
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            full_text = "\n".join(text_content)
            
            # Cache the content
            self.docx_cache[filename] = full_text
            
            print(f"✅ Successfully processed {filename} ({len(full_text)} characters)")
            return full_text
            
        except requests.RequestException as e:
            print(f"❌ Failed to download {filename}: {e}")
            return ""
        except Exception as e:
            print(f"❌ Failed to process {filename}: {e}")
            return ""

    def search_docx_files(self, question: str) -> str:
        """Search through all DOCX files and compile context"""
        if not DOCX_AVAILABLE:
            return ""
        
        print("📄 Searching DOCX files...")
        
        context_parts = []
        question_lower = question.lower()
        
        for filename in self.docx_files:
            content = self.download_docx_file(filename)
            if content:
                # Simple relevance check - if question keywords appear in content
                content_lower = content.lower()
                
                # Check for keyword matches
                question_words = question_lower.split()
                matches = sum(1 for word in question_words if len(word) > 2 and word in content_lower)
                
                if matches > 0:
                    # Include relevant sections (first 1000 chars to avoid token limits)
                    preview = content[:1000] + "..." if len(content) > 1000 else content
                    context_parts.append(f"Από αρχείο {filename}:\n{preview}")
                    print(f"✅ Found relevant content in {filename}")
        
        if context_parts:
            return "\n\n".join(context_parts)
        else:
            print("⚠️ No relevant DOCX content found")
            return ""

    def get_ai_response_with_docx(self, user_message: str) -> Tuple[str, bool]:
        """Get AI response using DOCX files as context"""
        if not self.groq_client:
            return "", False
        
        try:
            # Get DOCX context
            docx_context = self.search_docx_files(user_message)
            
            if not docx_context:
                return "", False
            
            # Prepare the full prompt
            full_prompt = f"""Βάσει των παρακάτω εγγράφων πρακτικής άσκησης:

{docx_context}

Ερώτηση φοιτητή: {user_message}

Χρησιμοποίησε ΜΟΝΟ τις πληροφορίες από τα έγγραφα για να απαντήσεις."""

            # Call Groq API
            chat_completion = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": full_prompt}
                ],
                model="llama-3.1-8b-instant",
                temperature=0.1,
                max_tokens=800,
                top_p=0.9,
                stream=False
            )

            response = chat_completion.choices[0].message.content
            
            # Έλεγχος για μη-ελληνικούς χαρακτήρες
            if response and any(ord(char) > 1500 and ord(char) not in range(0x0370, 0x03FF) for char in response):
                print("⚠️ Detected non-Greek characters in response, using fallback")
                return "", False
            
            return response, True
            
        except Exception as e:
            print(f"❌ DOCX AI Error: {e}")
            return "", False

    def calculate_similarity(self, question: str, qa_entry: Dict) -> float:
        """Calculate similarity between question and QA entry"""
        question_lower = question.lower()
        
        # Check if any keyword matches
        keyword_matches = sum(1 for keyword in qa_entry.get('keywords', []) 
                            if keyword.lower() in question_lower)
        
        # Check title similarity
        title_words = qa_entry['question'].lower().split()
        title_matches = sum(1 for word in title_words if word in question_lower)
        
        # Combine scores
        total_score = (keyword_matches * 2 + title_matches) / max(len(qa_entry.get('keywords', [])) + len(title_words), 1)
        return min(total_score, 1.0)

    def get_ai_response(self, user_message: str, context: str) -> Tuple[str, bool]:
        """Get response from Groq AI with context"""
        if not self.groq_client:
            return "", False
        
        try:
            # Prepare the full prompt with context
            full_prompt = f"""Context πληροφοριών:
{context}

Ερώτηση φοιτητή: {user_message}

Χρησιμοποίησε ΜΟΝΟ τις πληροφορίες από το context για να απαντήσεις."""

            # Call Groq API
            chat_completion = self.groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": user_message}
                ],
                model="llama-3.1-8b-instant",
                temperature=0.1,  # Χαμηλότερο για πιο συνεπείς απαντήσεις
                max_tokens=800,
                top_p=0.9,        # Πιο συντηρητικό για σταθερότητα
                stream=False
            )

            response = chat_completion.choices[0].message.content
            
            # Έλεγχος για μη-ελληνικούς χαρακτήρες
            if response and any(ord(char) > 1500 and ord(char) not in range(0x0370, 0x03FF) for char in response):
                print("⚠️ Detected non-Greek characters in response, using fallback")
                return "", False
            
            return response, True
            
        except Exception as e:
            print(f"❌ Groq API Error: {e}")
            return "", False

    def get_fallback_response(self, question: str) -> Tuple[str, bool]:
        """Fallback response system - returns (response, found_exact_match)"""
        if not self.qa_data:
            return "Δεν υπάρχουν διαθέσιμα δεδομένα. Επικοινωνήστε με τον Γεώργιο Σοφιανίδη: gsofianidis@mitropolitiko.edu.gr", False

        # Find best match
        best_match = max(self.qa_data, key=lambda x: self.calculate_similarity(question, x))
        similarity = self.calculate_similarity(question, best_match)

        if similarity > 0.2:
            return best_match['answer'], True
        else:
            return f"""Δεν βρέθηκε συγκεκριμένη απάντηση για αυτή την ερώτηση.

Προτεινόμενες ενέργειες:
• Αναδιατυπώστε την ερώτηση
• Επιλέξτε από τις συχνές ερωτήσεις στο αριστερό μενού
• Επικοινωνήστε με τον Γεώργιο Σοφιανίδη: gsofianidis@mitropolitiko.edu.gr""", False

    def get_response(self, question: str) -> str:
        """Get chatbot response - JSON FIRST, then DOCX AI, then JSON fallback"""
        if not self.qa_data:
            return "Δεν υπάρχουν διαθέσιμα δεδομένα γνώσης."
        
        # Step 1: Try JSON fallback FIRST
        json_response, found_exact_match = self.get_fallback_response(question)
        
        if found_exact_match:
            print("✅ Found exact match in JSON data")
            return json_response
        
        # Step 2: Try DOCX AI search
        print("📄 No good JSON match, trying DOCX AI search...")
        
        if self.groq_client and DOCX_AVAILABLE:
            docx_response, success = self.get_ai_response_with_docx(question)
            if success and docx_response.strip():
                print("✅ DOCX AI response successful")
                return docx_response
        
        # Step 3: Try regular AI with JSON context (fallback)
        print("🤖 DOCX search failed, trying regular AI...")
        
        if self.groq_client:
            # Find relevant context for AI
            matches = sorted(self.qa_data, 
                            key=lambda x: self.calculate_similarity(question, x), 
                            reverse=True)
            
            # Prepare context from top matches
            context_parts = []
            for match in matches[:3]:
                if self.calculate_similarity(question, match) > 0.1:
                    context_parts.append(f"Q: {match['question']}\nA: {match['answer']}")
            
            context = "\n\n".join(context_parts) if context_parts else ""
            
            if context:
                ai_response, success = self.get_ai_response(question, context)
                if success and ai_response.strip():
                    print("✅ Regular AI response successful")
                    return ai_response
        
        # Step 4: Final fallback to JSON (even if low similarity)
        print("📋 Using JSON fallback response")
        return json_response

def initialize_qa_file():
    """Create initial qa_data.json if it doesn't exist (fallback for development)"""
    if not os.path.exists("qa_data.json"):
        print("📄 qa_data.json not found. Please create it with the full 39 entries.")
        print("💡 Place the complete JSON file in the same directory as this script.")
        return False
    return True

def main():
    """Main Streamlit application - Git-first content management"""
    
    # CSS Styling
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #1f4e79 0%, #2980b9 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 2rem;
    }
    
    .header-content {
        flex: 1;
    }
    
    .header-logo {
        max-height: 80px;
        max-width: 120px;
        object-fit: contain;
    }
    
    .logo-container {
        display: flex;
        align-items: center;
        margin-bottom: 2rem;
        padding: 1rem 0;
        border-bottom: 1px solid #e8f4f8;
    }
    
    .user-message {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border: 1px solid #dee2e6;
        border-left: 4px solid #007bff;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #333;
    }
    
    .ai-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-left: 4px solid #4caf50;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        box-shadow: 0 4px 8px rgba(0,0,0,0.15);
    }
    
    .ai-message a {
        color: #ffeb3b !important;
        text-decoration: underline !important;
        font-weight: bold !important;
    }
    
    .ai-message a:hover {
        color: #fff9c4 !important;
        text-decoration: underline !important;
    }
    
    .info-card {
        background: white;
        border: 1px solid #e8f4f8;
        border-radius: 10px;
        padding: 1.5rem;
        margin: 0.5rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    
    .info-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 16px rgba(0,0,0,0.12);
    }
    
    .api-status {
        position: fixed;
        top: 20px;
        right: 20px;
        background: #28a745;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-size: 0.9rem;
        z-index: 1000;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .quick-stats {
        display: flex;
        justify-content: space-around;
        margin: 2rem 0;
        gap: 1rem;
    }
    
    .stat-item {
        text-align: center;
        background: white;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #e8f4f8;
        flex: 1;
    }
    
    .chat-container {
        background: white;
        border-radius: 10px;
        padding: 2rem;
        margin-top: 2rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        border: 1px solid #e8f4f8;
    }
    
    .stTextInput > div > div > input {
        border-radius: 20px;
        border: 2px solid #e8f4f8;
        padding: 0.8rem 1.2rem;
    }
    
    .stButton > button {
        background: linear-gradient(90deg, #1f4e79 0%, #2980b9 100%);
        color: white;
        border: none;
        border-radius: 20px;
        padding: 0.6rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(31, 78, 121, 0.3);
    }
    </style>
    """, unsafe_allow_html=True)

    # Header with Logo
    logo_url = "https://raw.githubusercontent.com/GiorgosBouh/chatbot.placement/main/MK_LOGO_SEO_1200x630.png"
    
    st.markdown(f"""
    <div class="main-header">
        <img src="{logo_url}" alt="Μητροπολιτικό Κολλέγιο" class="header-logo">
        <div class="header-content">
            <h1>Πρακτική Άσκηση</h1>
            <h3>Μητροπολιτικό Κολλέγιο - Τμήμα Προπονητικής & Φυσικής Αγωγής</h3>
            <p><em>Εξειδικευμένος AI Assistant για υποστήριξη φοιτητών</em></p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Initialize session state
    if 'messages' not in st.session_state:
        st.session_state.messages = []

    if 'chatbot' not in st.session_state:
        # Get Groq API key from secrets or environment
        groq_api_key = None
        try:
            groq_api_key = st.secrets.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY")
        except:
            pass
        st.session_state.chatbot = InternshipChatbot(groq_api_key)
    else:
        # Refresh data if cache was cleared
        current_data_count = len(st.session_state.chatbot.qa_data)
        st.session_state.chatbot.qa_data = st.session_state.chatbot.load_qa_data()
        new_data_count = len(st.session_state.chatbot.qa_data)
        
        if new_data_count != current_data_count:
            st.toast(f"📊 Δεδομένα ενημερώθηκαν: {new_data_count} ερωτήσεις")

    # Quick info cards
    st.markdown("### 📊 Σημαντικές Πληροφορίες")
    
    quick_col1, quick_col2, quick_col3 = st.columns(3)
    
    with quick_col1:
        st.markdown("""
        <div class="info-card" style="text-align: center;">
            <h4 style="color: #1f4e79; margin-bottom: 0.5rem;">📅 Απαιτούμενες Ώρες</h4>
            <p style="font-size: 1.2rem; font-weight: 600; color: #28a745; margin: 0;">240 ώρες</p>
            <small style="color: #6c757d;">Προθεσμία: 30 Μαϊου</small>
        </div>
        """, unsafe_allow_html=True)

    with quick_col2:
        st.markdown("""
        <div class="info-card" style="text-align: center;">
            <h4 style="color: #1f4e79; margin-bottom: 0.5rem;">📋 Παράδοση Συμβάσεων</h4>
            <p style="font-size: 1.2rem; font-weight: 600; color: #ffc107; margin: 0;">Moodle Platform</p>
            <small style="color: #6c757d;">Προθεσμία: 15 Οκτωβρίου</small>
        </div>
        """, unsafe_allow_html=True)

    with quick_col3:
        st.markdown("""
        <div class="info-card" style="text-align: center;">
            <h4 style="color: #1f4e79; margin-bottom: 0.5rem;">⏰ Επιτρεπόμενο Ωράριο</h4>
            <p style="font-size: 1.2rem; font-weight: 600; color: #17a2b8; margin: 0;">Δευτέρα-Σάββατο</p>
            <small style="color: #6c757d;">Μέχρι 8 ώρες ανά ημέρα</small>
        </div>
        """, unsafe_allow_html=True)

    # API Status
    if st.session_state.chatbot.groq_client:
        if DOCX_AVAILABLE:
            st.markdown('<div class="api-status">📋 JSON + DOCX Mode</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="api-status">📋 JSON First Mode</div>', unsafe_allow_html=True)
        
    # Επαγγελματική ενδειξη για sidebar
    st.markdown("""
    <div style="background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 4px; padding: 0.6rem; margin-bottom: 1.5rem; text-align: center; font-size: 0.9rem;">
        <strong>Πληροφορίες:</strong> Χρησιμοποιήστε το αριστερό μενού για συχνές ερωτήσεις και επικοινωνία 👈<br>
        <small>🔄 Προτεραιότητα: JSON → DOCX → AI</small>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown("## 📞 Επικοινωνία")
        
        st.markdown("""
        **Υπεύθυνος Πρακτικής Άσκησης**  
        **Γεώργιος Σοφιανίδης**  
        📧 gsofianidis@mitropolitiko.edu.gr
        
        **Τεχνική Υποστήριξη**  
        **Γεώργιος Μπουχουράς**  
        📧 gbouchouras@mitropolitiko.edu.gr
        """)

        st.markdown("---")

        # Συχνές ερωτήσεις
        st.markdown("## 🔄 Συχνές Ερωτήσεις")
        
        # Group questions by category
        categories = {}
        for qa in st.session_state.chatbot.qa_data:
            cat = qa.get('category', 'Άλλα')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(qa)

        for category, questions in categories.items():
            if st.expander(f"📂 {category}"):
                for qa in questions:
                    if st.button(qa['question'], key=f"faq_{qa['id']}", use_container_width=True):
                        # Add to chat
                        st.session_state.messages.append({"role": "user", "content": qa['question']})
                        st.session_state.messages.append({"role": "assistant", "content": qa['answer']})
                        st.rerun()

        st.markdown("---")

        # AI Status
        if st.session_state.chatbot.groq_client:
            if DOCX_AVAILABLE:
                st.success("📋 JSON + DOCX Mode")
                st.info("AI ψάχνει σε επίσημα έγγραφα")
            else:
                st.success("📋 JSON First Mode")
                st.warning("DOCX search απενεργοποιημένο")
        else:
            st.warning("📚 JSON Only Mode")
            if GROQ_AVAILABLE:
                st.info("Για AI+DOCX, χρειάζεται Groq API key")
            else:
                st.error("Groq library δεν είναι διαθέσιμη")

        st.markdown("---")

        if st.button("🗑️ Νέα Συνομιλία", use_container_width=True):
            st.session_state.messages = []
            st.rerun()

        # Enhanced Technical Information
        if st.checkbox("🔧 Τεχνικές Πληροφορίες"):
            st.markdown("**Για τεχνικά προβλήματα:**")
            st.markdown("📧 gbouchouras@mitropolitiko.edu.gr")
            
            # Enhanced debugging info
            st.write("**System Status:**")
            st.write("• Response Priority: JSON → DOCX → AI")
            st.write("• Groq Available:", GROQ_AVAILABLE)
            st.write("• Groq Client:", st.session_state.chatbot.groq_client is not None)
            st.write("• DOCX Available:", DOCX_AVAILABLE)
            st.write("• QA Data Count:", len(st.session_state.chatbot.qa_data))
            
            # DOCX Status
            if DOCX_AVAILABLE:
                st.write("**DOCX Files:**")
                for filename in st.session_state.chatbot.docx_files:
                    cached = "📋" if filename in st.session_state.chatbot.docx_cache else "⏳"
                    st.write(f"• {cached} {filename}")
                
                if st.session_state.chatbot.docx_cache:
                    total_chars = sum(len(content) for content in st.session_state.chatbot.docx_cache.values())
                    st.info(f"📊 Cached DOCX content: {total_chars:,} characters")
            else:
                st.error("📄 DOCX processing disabled - install python-docx")
            
            # File status
            qa_file_exists = os.path.exists("qa_data.json")
            st.write("• qa_data.json exists:", qa_file_exists)
            
            if qa_file_exists:
                try:
                    with open("qa_data.json", 'r', encoding='utf-8') as f:
                        file_data = json.load(f)
                    st.success(f"📄 External JSON: {len(file_data)} entries loaded")
                    st.write(f"• Entry IDs: {[d['id'] for d in file_data[:5]]}")
                    if len(file_data) > 5:
                        st.write(f"• ... and {len(file_data)-5} more")
                    
                    # File info
                    file_size = os.path.getsize("qa_data.json")
                    mtime = os.path.getmtime("qa_data.json")
                    last_modified = datetime.datetime.fromtimestamp(mtime).strftime("%d/%m/%Y %H:%M")
                    st.info(f"📊 File size: {file_size:,} bytes")
                    st.info(f"🕒 Last modified: {last_modified}")
                    
                except Exception as e:
                    st.error(f"❌ JSON Error: {e}")
            else:
                st.warning("📋 Using fallback data")
                st.error("💡 Create qa_data.json with 39 entries!")
            
            # Directory info
            st.write("**File System:**")
            st.write("• Current dir:", os.getcwd())
            files = [f for f in os.listdir('.') if f.endswith('.json')]
            st.write("• JSON files:", files if files else "None found")
            
            # Categories info
            if st.session_state.chatbot.qa_data:
                categories_count = {}
                for qa in st.session_state.chatbot.qa_data:
                    cat = qa.get('category', 'Unknown')
                    categories_count[cat] = categories_count.get(cat, 0) + 1
                
                st.write("**Categories:**")
                for cat, count in categories_count.items():
                    st.write(f"• {cat}: {count}")

    # Chat interface
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    st.markdown("### 💬 Κάντε την ερώτησή σας")

    # Display chat messages
    for message in st.session_state.messages:
        if message["role"] == "user":
            st.markdown(f'<div class="user-message"><strong>Εσείς:</strong> {message["content"]}</div>', unsafe_allow_html=True)
        else:
            # Convert markdown to HTML for better display
            content = message["content"].replace('\n', '<br>')
            st.markdown(f'<div class="ai-message"><strong>🤖 Assistant:</strong><br><br>{content}</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # Chat input - moved outside container for better functionality
    user_input = st.chat_input("Γράψτε την ερώτησή σας εδώ...")
    
    if user_input:
        # Add user message
        st.session_state.messages.append({"role": "user", "content": user_input})
        
        # Get chatbot response
        with st.spinner("Σκέφτομαι..."):
            try:
                response = st.session_state.chatbot.get_response(user_input)
            except Exception as e:
                response = f"Συγγνώμη, παρουσιάστηκε σφάλμα: {str(e)}"
                st.error(f"Σφάλμα: {e}")
        
        # Add assistant response
        st.session_state.messages.append({"role": "assistant", "content": response})
        
        # Rerun to display new messages
        st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #6c757d; padding: 1rem;">
        <small>
            🎓 <strong>Μητροπολιτικό Κολλέγιο Θεσσαλονίκης</strong> | 
            Τμήμα Προπονητικής & Φυσικής Αγωγής<br>
            <em>JSON-First + DOCX AI Assistant</em>
        </small>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()